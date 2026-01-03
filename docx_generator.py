from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_docx(
    report_title,
    brand_title,
    date_str,
    keywords,
    articles,
    out_path
):
    doc = Document()

    # ---- HEADER ----
    doc.add_heading(report_title, level=1)
    doc.add_paragraph(f"Brand / Subject: {brand_title}")
    doc.add_paragraph(f"Report Date: {date_str}")

    if keywords:
        doc.add_paragraph(f"Keywords Searched: {', '.join(keywords)}")

    doc.add_paragraph("")

    # ---- SUMMARY ----
    total_articles = len(articles)
    flagged_articles = [
        a for a in articles
        if a.get("author_flag") != "Named Author" or a.get("marked_ai")
    ]

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(f"Total articles included: {total_articles}")
    doc.add_paragraph(
        f"Articles with authorship concerns: {len(flagged_articles)}"
    )

    doc.add_paragraph("")

    # ---- TABLE ----
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Article"
    hdr[2].text = "Source"
    hdr[3].text = "Author"
    hdr[4].text = "Authorship Status"

    for a in articles:
        row = table.add_row().cells

        row[0].text = a.get("date", "")

        p = row[1].paragraphs[0]
        add_hyperlink(
            p,
            a.get("title", ""),
            a.get("url", "")
        )

        row[2].text = a.get("source", "")
        row[3].text = a.get("author", "") or "—"

        status = a.get("author_flag", "Unknown")
        if a.get("marked_ai"):
            status += " (Manually flagged)"

        row[4].text = status

    doc.save(out_path)
