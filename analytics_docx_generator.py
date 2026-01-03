from docx import Document
from docx.shared import Inches
import altair as alt
import pandas as pd
import os
import tempfile


# --------------------------------------------------
# Utility: save altair chart as image
# --------------------------------------------------
def save_chart(chart, filename):
    chart.save(filename, scale_factor=2.0)


# --------------------------------------------------
# MAIN GENERATOR
# --------------------------------------------------
def generate_analytics_docx(df: pd.DataFrame, meta: dict, out_path: str):
    doc = Document()

    # ---------------- HEADER ----------------
    doc.add_heading("Media Analytics Report", level=1)

    doc.add_paragraph(
        f"Analysis Period: {meta['from']} → {meta['to']}"
    )
    doc.add_paragraph(
        f"Keywords: {', '.join(meta['keywords'])}"
    )

    doc.add_paragraph("")

    # ---------------- KPIs ----------------
    total = len(df)
    news = len(df[df["content_type"] == "News"])
    blogs = len(df[df["content_type"] == "Blog"])
    sources = df["source"].nunique()

    risk = (
        round(
            100 * len(
                df[(df["author_flag"] != "Named Author") | (df["marked_ai"])]
            ) / total,
            1
        )
        if total else 0
    )

    doc.add_heading("Key Metrics", level=2)
    doc.add_paragraph(f"Total Articles: {total}")
    doc.add_paragraph(f"News / Blogs: {news} / {blogs}")
    doc.add_paragraph(f"Unique Sources: {sources}")
    doc.add_paragraph(f"Authorship Risk: {risk}%")

    doc.add_page_break()

    # temp dir for images
    tmp = tempfile.mkdtemp()
    # Normalize dates for Altair / JSON safety
    if "date" in df.columns:
        df["date"] = pd.to_datetime(df["date"]).dt.strftime("%Y-%m-%d")

    # ==================================================
    # 📈 Mentions Over Time
    # ==================================================
    daily = (
        df.groupby("date")
        .size()
        .reset_index(name="count")
        .sort_values("date")
    )

    # 🔥 CRITICAL FIX
    daily["date"] = pd.to_datetime(daily["date"]).dt.strftime("%Y-%m-%d")

    mentions_chart = alt.Chart(daily).mark_line(point=True).encode(
        x=alt.X("date:T", title="Date"),
        y=alt.Y("count:Q", title="Articles"),
        tooltip=["date", "count"]
    )

    path = os.path.join(tmp, "mentions_over_time.png")
    save_chart(mentions_chart, path)

    doc.add_heading("Mentions Over Time", level=2)
    doc.add_picture(path, width=Inches(6))

    # ==================================================
    # 🗞️ Top Publications
    # ==================================================
    top_sources = (
        df["source"]
        .fillna("Unknown")
        .value_counts()
        .head(10)
        .reset_index()
    )
    top_sources.columns = ["Source", "Count"]

    sources_chart = alt.Chart(top_sources).mark_bar().encode(
        y=alt.Y("Source:N", sort="-x", title="Publication"),
        x=alt.X("Count:Q", title="Articles"),
        tooltip=["Source", "Count"]
    )

    path = os.path.join(tmp, "top_sources.png")
    save_chart(sources_chart, path)

    doc.add_heading("Top Publications", level=2)
    doc.add_picture(path, width=Inches(6))

    # ==================================================
    # 📰 News vs Blogs
    # ==================================================
    nb = df["content_type"].value_counts().reset_index()
    nb.columns = ["Type", "Count"]

    nb_chart = alt.Chart(nb).mark_arc().encode(
        theta="Count:Q",
        color="Type:N",
        tooltip=["Type", "Count"]
    )

    path = os.path.join(tmp, "news_vs_blogs.png")
    save_chart(nb_chart, path)

    doc.add_heading("News vs Blogs", level=2)
    doc.add_picture(path, width=Inches(4.5))

    # ==================================================
    # 🔑 Keyword Performance
    # ==================================================
    keyword_hits = {
        kw: df["title"].astype(str).str.contains(kw, case=False, na=False).sum()
        for kw in meta["keywords"]
    }

    kw_df = pd.DataFrame(
        list(keyword_hits.items()),
        columns=["Keyword", "Mentions"]
    )

    kw_chart = alt.Chart(kw_df).mark_bar().encode(
        x=alt.X("Keyword:N", title="Keyword"),
        y=alt.Y("Mentions:Q", title="Mentions"),
        tooltip=["Keyword", "Mentions"]
    )

    path = os.path.join(tmp, "keyword_performance.png")
    save_chart(kw_chart, path)

    doc.add_heading("Keyword Performance", level=2)
    doc.add_picture(path, width=Inches(6))

    # ==================================================
    # ✍️ Authorship Quality
    # ==================================================
    auth = df.apply(
        lambda r: "AI / Flagged" if r["marked_ai"] else r["author_flag"],
        axis=1
    ).value_counts().reset_index()
    auth.columns = ["Category", "Count"]

    auth_chart = alt.Chart(auth).mark_arc().encode(
        theta="Count:Q",
        color="Category:N",
        tooltip=["Category", "Count"]
    )

    path = os.path.join(tmp, "authorship_quality.png")
    save_chart(auth_chart, path)

    doc.add_heading("Authorship Quality", level=2)
    doc.add_picture(path, width=Inches(4.5))

    # ---------------- SAVE ----------------
    doc.save(out_path)
